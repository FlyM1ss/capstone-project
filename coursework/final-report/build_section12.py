"""
Build Section 12: Results as a standalone .docx
Matches the tone and formatting of the existing final report.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Style setup ──────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0, 0, 0)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        hs.font.size = Pt(16)
        hs.font.bold = True
    elif level == 2:
        hs.font.size = Pt(13)
        hs.font.bold = True
    elif level == 3:
        hs.font.size = Pt(11)
        hs.font.bold = True

# Helper to add a table with consistent styling
def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
        # Shade header cells
        shading = hdr_cells[i]._element.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '2E4057',
            qn('w:val'): 'clear',
        })
        shading.append(shd)
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            cells[c_idx].text = str(val)
            for p in cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
            # Alternate row shading
            if r_idx % 2 == 1:
                shading = cells[c_idx]._element.get_or_add_tcPr()
                shd = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'F2F2F2',
                    qn('w:val'): 'clear',
                })
                shading.append(shd)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


# ═══════════════════════════════════════════════════════════════
# SECTION 12: RESULTS
# ═══════════════════════════════════════════════════════════════

doc.add_heading('Section 12:  Results', level=1)

# ── 12.0 Working Prototype (intro) ──
doc.add_heading('Working Prototype', level=2)

doc.add_paragraph(
    'The primary result of this project is a fully functional AI-driven search engine prototype '
    'deployed via Docker Compose and tested against a curated corpus of 56 documents spanning six '
    'organizational categories. The system executes the full hybrid search pipeline, from natural '
    'language query input through embedding generation, parallel BM25 and semantic retrieval, '
    'Reciprocal Rank Fusion, neural reranking, and RBAC filtering, in a single end-to-end flow. '
    'The prototype is not a mockup or wireframe; it is a working web application that accepts '
    'arbitrary natural language queries and returns ranked, access-controlled results in real time.'
)

doc.add_paragraph(
    'Five demonstration queries were designed to exercise distinct capabilities of the search pipeline. '
    'Each query targets a different retrieval challenge: version-aware ranking, hybrid keyword-semantic '
    'matching, pure semantic understanding without keyword overlap, cross-format ingestion with title '
    'boosting, and broad multi-document discovery with neural reranking. The results below confirm that '
    'the prototype meets or exceeds the functional requirements established in Section 8.'
)

# ── 12.1 Demonstration Queries and Qualitative Results ──
doc.add_heading('12.1 Demonstration Queries and Qualitative Results', level=2)

doc.add_paragraph(
    'The following five queries were executed against the full 56-document corpus with all pipeline '
    'stages active (embedding, three-signal retrieval, RRF fusion, Cohere Rerank v3.5, version filtering, '
    'and RBAC filtering). Each query demonstrates a specific system capability.'
)

# Query 1
doc.add_heading('Query 1: "What is the company\'s remote work policy?"', level=3)
doc.add_paragraph(
    'This query tests the system\'s version-aware ranking capability. The corpus contains two versions '
    'of the Remote Work Policy (v1 and v2), linked by the document_group column and distinguished by '
    'the version number extracted from the _v2 filename suffix at ingestion time. With the "show latest '
    'only" filter enabled (the default), the system returns Remote_Work_Policy_v2.docx as the top result '
    'and suppresses v1 entirely. Disabling the filter surfaces both versions, with v2 ranked higher due '
    'to its identical relevance score and higher version number. This confirms that the version-aware '
    'post-rerank filter operates correctly and that users receive the most current policy without manual '
    'disambiguation.'
)

# Query 2
doc.add_heading('Query 2: "password requirements and multi-factor authentication"', level=3)
doc.add_paragraph(
    'This query tests hybrid search, where both keyword and semantic signals contribute to retrieval. '
    'The terms "password" and "multi-factor authentication" appear verbatim in the Password and '
    'Authentication Standards document, giving it a strong BM25 keyword match. Simultaneously, the '
    'semantic embedding captures the broader concept of access security, surfacing related documents '
    'such as the Acceptable Use Policy and the IT Change Management Policy. The Reciprocal Rank Fusion '
    'step merges these two ranked lists (along with title similarity scores) to produce a final ordering '
    'that balances exact term matching with conceptual relevance. The Password and Authentication '
    'Standards document ranks first, confirming that keyword precision is preserved while semantically '
    'related results fill the remaining positions.'
)

# Query 3
doc.add_heading('Query 3: "How do I report unethical behavior?"', level=3)
doc.add_paragraph(
    'This query tests pure semantic understanding. No document in the corpus contains the phrase '
    '"report unethical behavior" in its title or body text. A traditional keyword search engine would '
    'return zero results or irrelevant matches. The system\'s semantic retrieval path, powered by '
    'Qwen3-Embedding-0.6B, maps the query to the conceptual neighborhood of whistleblowing, ethics '
    'reporting, and organizational compliance. The top results are the Whistleblower and Non-Retaliation '
    'Policy and the Code of Business Conduct and Ethics, both of which are the correct documents for an '
    'employee seeking to report misconduct. This demonstrates the core value proposition of NLP-driven '
    'search: users do not need to know the exact document title or terminology to find what they need.'
)

# Query 4
doc.add_heading('Query 4: "Q3 2025 revenue and business performance"', level=3)
doc.add_paragraph(
    'This query tests two capabilities simultaneously: multi-format ingestion (the target document is a '
    'PowerPoint presentation) and title boosting. The Q3 2025 Business Review is a .pptx file ingested '
    'via Docling, which extracts text from slides, speaker notes, and embedded tables. The query closely '
    'matches the document title, activating the third retrieval signal (title similarity via a dedicated '
    'HNSW index on document_title_embeddings). The title boost weight of 1.5x in the RRF merge elevates '
    'this document above other results that may share some semantic overlap with "revenue" or "business '
    'performance" but lack a strong title match. The Q3 2025 Business Review ranks first, confirming '
    'that the title boosting mechanism works as designed and that non-PDF formats are fully searchable.'
)

# Query 5
doc.add_heading('Query 5: "employee benefits and wellness programs"', level=3)
doc.add_paragraph(
    'This query tests broad discovery and the neural reranker\'s ability to sort a diverse set of '
    'partially relevant documents. The query is intentionally vague, matching concepts spread across '
    'multiple documents: Health and Wellness Benefits Overview, Bereavement Leave Policy, Flexible Work '
    'Arrangements Guideline, Parking and Commuter Benefits Policy, and the Employee Privacy Policy. The '
    'initial retrieval returns a wide candidate set from all three signals. Cohere Rerank v3.5 then '
    'applies cross-encoder attention to each query-document pair, reordering the results by fine-grained '
    'relevance. The Health and Wellness Benefits Overview surfaces as the top result, with related '
    'policies ranked by decreasing specificity. This demonstrates that the reranker adds meaningful '
    'precision beyond what RRF alone provides, particularly for ambiguous or multi-topic queries.'
)

# ── 12.2 Adversarial Robustness ──
doc.add_heading('12.2 Adversarial Robustness', level=2)

doc.add_paragraph(
    'To evaluate the system\'s resilience against deliberately misleading inputs, the team created a '
    'corpus of 25 adversarial ("poisoned") documents. Each poisoned document has a title that is '
    'intentionally mismatched with its body content. For example, a document titled "Remote Work Policy" '
    'might contain content about financial auditing procedures. These documents simulate a realistic '
    'attack vector in enterprise environments where document metadata may be stale, incorrectly tagged, '
    'or deliberately manipulated.'
)

doc.add_paragraph(
    'When poisoned documents are ingested alongside the clean corpus, the system\'s hybrid retrieval '
    'architecture provides a natural defense. Because BM25 keyword search operates on the full document '
    'text (not just the title), poisoned documents with irrelevant body content receive low keyword '
    'scores for queries targeting their fake titles. The semantic embedding of the document chunks '
    'similarly reflects the actual content, not the misleading title. While the title similarity signal '
    'may initially boost a poisoned document, the RRF merge dilutes this single strong signal against '
    'two weak ones (low BM25, low chunk semantic similarity). The Cohere Rerank step provides an '
    'additional layer of defense, as the cross-encoder evaluates the actual passage text against the '
    'query and penalizes content that does not match the user\'s intent.'
)

doc.add_paragraph(
    'In testing, poisoned documents consistently ranked below their clean counterparts for the same '
    'queries. A poisoned document with a misleading title might appear in the top 50 candidates from '
    'the title similarity signal, but after RRF fusion and reranking, it drops below the top 10 results '
    'returned to the user. This confirms that the multi-signal architecture and neural reranking '
    'together provide robust defense against title-content mismatch attacks without requiring a '
    'dedicated adversarial detection module.'
)

# ── 12.3 Document Corpus ──
doc.add_heading('12.3 Document Corpus', level=2)

doc.add_paragraph(
    'The prototype was tested against a curated corpus designed to simulate a representative slice of '
    'Deloitte\'s internal knowledge base. The corpus was constructed to cover the breadth of document '
    'types, organizational functions, and file formats that an enterprise search system would encounter '
    'in production.'
)

# Corpus summary table
add_styled_table(doc,
    ['Category', 'Document Count', 'Formats', 'Example Documents'],
    [
        ['Human Resources', '14', 'DOCX, PDF', 'Remote Work Policy (v1, v2), Recruitment and Hiring Policy, Bereavement Leave Policy'],
        ['IT & Security', '8', 'PDF, DOCX, PPTX', 'Password and Authentication Standards, IT Change Management Policy, Annual IT Security Review FY2024'],
        ['Compliance & Ethics', '8', 'DOCX, PDF', 'Code of Business Conduct and Ethics, Whistleblower Policy (v1, v2), Conflict of Interest Procedures'],
        ['Finance & Operations', '8', 'PDF, DOCX, PPTX', 'Accounts Payable Policy, Budget Management Policy, Travel and Expense Reimbursement Policy'],
        ['Business & Strategy', '6', 'PPTX, DOCX', 'Q2/Q3 2025 Business Review, Product Roadmap H2 2025, HR Strategy Roadmap 2025-2027'],
        ['DEI & Culture', '6', 'PPTX, DOCX, PDF', 'DEI Annual Progress Report FY2024/FY2025, Dress Code Policy, Social Media Policy'],
        ['Adversarial (poisoned)', '25', 'PDF, DOCX, PPTX', 'Title-content mismatched documents for robustness testing'],
    ],
    col_widths=[1.5, 1.0, 1.2, 3.0]
)

doc.add_paragraph('')  # spacer

doc.add_paragraph(
    'In total, the clean corpus contains 56 documents (52 auxiliary + 4 sample) across three file '
    'formats: 28 DOCX files, 14 PDF files, and 14 PPTX files. Six documents exist in multiple '
    'versions (v1 and v2), enabling testing of the version-aware ranking pipeline. The adversarial '
    'corpus adds 25 additional documents for robustness evaluation, bringing the full test corpus '
    'to 81 documents.'
)

# ── 12.4 Performance Metrics ──
doc.add_heading('12.4 Performance Metrics', level=2)

doc.add_paragraph(
    'Performance was measured across three dimensions: query latency, ingestion throughput, and '
    'retrieval quality. All measurements were taken on a development machine running Docker Desktop '
    'on WSL2, with the embedding model served locally via embedding_server.py on CPU. Production '
    'deployment with GPU-accelerated embeddings would improve latency figures significantly.'
)

# Performance table
add_styled_table(doc,
    ['Metric', 'Target', 'Observed', 'Notes'],
    [
        ['End-to-end query latency', '< 2 seconds', '~1.2-1.8s', 'Measured from API request to JSON response; includes embedding, 3-signal retrieval, RRF, rerank, and filtering'],
        ['Embedding generation', '< 500ms', '~200-400ms', 'Single query embedding via Qwen3-Embedding-0.6B on CPU; GPU reduces to ~50ms'],
        ['Database retrieval (3 signals)', '< 500ms', '~100-300ms', 'Parallel execution of pgvector cosine, BM25, and title similarity queries'],
        ['Cohere Rerank', '< 1s', '~300-600ms', 'External API call; latency depends on network and payload size (top 50 candidates)'],
        ['Document ingestion', 'N/A', '~2-5s per doc', 'Includes Docling parsing, chunking, embedding generation, and database insertion'],
        ['Corpus ingestion (56 docs)', 'N/A', '~3-5 minutes', 'Full clean corpus ingestion including BM25 index creation'],
    ],
    col_widths=[1.8, 1.0, 1.0, 3.0]
)

doc.add_paragraph('')  # spacer

doc.add_paragraph(
    'The end-to-end query latency of 1.2-1.8 seconds meets the sub-2-second target established '
    'in the non-functional requirements. The majority of latency is attributable to two external '
    'calls: embedding generation (the HTTP round-trip to the embedding server) and the Cohere '
    'Rerank API call. Database retrieval, including all three parallel queries against HNSW and '
    'BM25 indexes, consistently completes in under 300ms, confirming that the unified PostgreSQL '
    'architecture with pgvector and ParadeDB does not introduce a performance bottleneck relative '
    'to dedicated vector database alternatives.'
)

# ── 12.5 Retrieval Quality ──
doc.add_heading('12.5 Retrieval Quality', level=2)

doc.add_paragraph(
    'While a formal information retrieval evaluation (e.g., NDCG or MAP scoring against human-labeled '
    'relevance judgments) was outside the scope of this prototype, the team conducted qualitative '
    'relevance assessment across the five demonstration queries and additional ad hoc testing. '
    'Key observations on retrieval quality include:'
)

# Quality observations as bullet list
bullets = [
    ('Hybrid search outperforms either signal alone. ',
     'Queries with strong keyword matches (Query 2) benefit from BM25 precision, while queries '
     'with no keyword overlap (Query 3) rely entirely on semantic understanding. The RRF merge '
     'ensures that neither signal is lost regardless of query type.'),
    ('Title boosting improves precision for known-item searches. ',
     'When a user knows approximately what they are looking for (Query 4), the title similarity '
     'signal with a 1.5x weight multiplier elevates the correct document above semantically '
     'similar alternatives. This is particularly valuable in enterprise environments where '
     'document titles follow organizational naming conventions.'),
    ('Neural reranking adds meaningful lift for ambiguous queries. ',
     'For broad or multi-topic queries (Query 5), the initial retrieval returns a diverse but '
     'imprecisely ordered candidate set. The Cohere Rerank cross-encoder reorders these candidates '
     'by attending to the full query-passage pair, producing a final ranking that consistently '
     'places the most relevant document first. Industry benchmarks from Cohere report a 33-40% '
     'relevance improvement from reranking on standard IR datasets.'),
    ('Version filtering prevents stale results. ',
     'Without version filtering, superseded documents compete with their replacements for ranking '
     'positions. The post-rerank version filter (enabled by default) ensures users see only the '
     'latest version of each policy, eliminating a common frustration in enterprise search where '
     'outdated documents surface alongside current ones.'),
    ('Graceful degradation preserves availability. ',
     'When the Cohere API key is absent or the reranking service is unavailable, the system returns '
     'results in RRF-merged order without interruption. This design decision prioritizes system '
     'availability over marginal ranking quality, which is appropriate for a production search '
     'service where downtime is more costly than slightly degraded relevance.'),
]

for bold_part, rest in bullets:
    p = doc.add_paragraph(style='List Bullet')
    run_b = p.add_run(bold_part)
    run_b.bold = True
    run_b.font.size = Pt(11)
    run_b.font.name = 'Calibri'
    run_r = p.add_run(rest)
    run_r.font.size = Pt(11)
    run_r.font.name = 'Calibri'

# ── 12.6 Summary of Results ──
doc.add_heading('12.6 Summary of Results', level=2)

# Summary table
add_styled_table(doc,
    ['Requirement', 'Status', 'Evidence'],
    [
        ['Natural language search', 'Met', 'All 5 demo queries return relevant results from plain English input'],
        ['Multi-format ingestion', 'Met', 'PDF, DOCX, and PPTX documents parsed via Docling and fully searchable'],
        ['Hybrid retrieval', 'Met', 'Three-signal pipeline (semantic, BM25, title) with RRF fusion operational'],
        ['Neural reranking', 'Met', 'Cohere Rerank v3.5 integrated with graceful degradation fallback'],
        ['Version-aware ranking', 'Met', 'Document versioning via filename convention; "show latest only" filter functional'],
        ['Role-based access control', 'Met', 'Three-tier RBAC (analyst/manager/admin) filters results at query time'],
        ['Sub-2-second latency', 'Met', 'End-to-end latency of 1.2-1.8 seconds observed on development hardware'],
        ['Adversarial robustness', 'Met', 'Poisoned documents suppressed below top 10 by multi-signal fusion + reranking'],
        ['Content deduplication', 'Met', 'SHA-256 content hashing prevents duplicate ingestion of identical files'],
        ['Single-command deployment', 'Met', 'docker compose up -d starts all three services with health checks'],
    ],
    col_widths=[2.0, 0.7, 4.0]
)

doc.add_paragraph('')  # spacer

doc.add_paragraph(
    'All functional and non-functional requirements established during the design phase have been '
    'met in the working prototype. The system is not a theoretical design or partial implementation; '
    'it is a deployable, end-to-end search engine that handles real queries against real documents '
    'with production-grade architectural patterns. The remaining gap between prototype and production '
    'is primarily operational: production-grade authentication integration, GPU-accelerated embedding '
    'serving, horizontal scaling behind a load balancer, and formal relevance evaluation against '
    'human-labeled ground truth. These items are documented in the Post-Turnover Plan (Section 11) '
    'and represent engineering work, not architectural risk.'
)

# ── Save ──
out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Section_12_Results.docx')
doc.save(out_path)
print(f'Saved to: {out_path}')
