# coursework/final-report/apply_edits.py
"""
Apply surgical polish patches to Final Report Draft.docx.

Usage:
    python apply_edits.py                 # run all patches -> Final Report Draft.final.docx
    python apply_edits.py --dry-run       # run patches against a temp copy, verify only
    python apply_edits.py --scan          # re-enumerate locators and print the values
    python apply_edits.py --only N        # run only patch N (1..5)
    python apply_edits.py --up-to N       # run patches 1..N
"""
from __future__ import annotations

import argparse
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import locators as L

HERE = Path(__file__).parent
SOURCE = HERE / "Final Report Draft.docx"
OUTPUT = HERE / "Final Report Draft.final.docx"


def scan(source: Path) -> None:
    doc = Document(source)
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")
    print(f"Inline shapes: {len(doc.inline_shapes)}")
    print("\nHeadings:")
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith("Heading") and p.text.strip():
            print(f"  {i:4d} [{p.style.name:10s}] {p.text[:80]}")
    print("\nTables:")
    for i, tbl in enumerate(doc.tables):
        r0 = tbl.rows[0]
        preview = " | ".join(c.text.strip()[:20] for c in r0.cells[:3])
        print(f"  {i:2d}: {len(tbl.rows)}x{len(tbl.columns)}  hdr: {preview}")


# ═════════════════════════════════════════════════════════════════════
# Patch 1: §8.10 screenshots
# ═════════════════════════════════════════════════════════════════════

SCREENSHOT_SPECS = [
    (
        "landing.png",
        "Figure 8.10.1: Landing page with natural language search input.",
        "The landing page accepts plain-English queries and surfaces the core "
        "value proposition above the fold.",
    ),
    (
        "search-results.png",
        "Figure 8.10.2: Search results with relevance scores, metadata, and filters.",
        "Each result card shows the document title, a scored relevance indicator, "
        "snippet highlights, and filter controls for role-based access and document type.",
    ),
    (
        "admin-upload.png",
        "Figure 8.10.3: Admin upload portal with drag-and-drop ingestion.",
        "Administrators can queue new documents through a drag-and-drop panel; "
        "ingestion status is reported inline.",
    ),
]


def patch_1_8_10_screenshots(docx_path: Path, screenshots: dict[str, Path]) -> None:
    doc = Document(docx_path)

    placeholder = doc.paragraphs[L.PARA_PLACEHOLDER_8_10]
    if "will be included in the final version" not in placeholder.text:
        raise RuntimeError(
            f"§8.10 placeholder not found at paragraph {L.PARA_PLACEHOLDER_8_10}. "
            f"Got: {placeholder.text[:120]!r}"
        )

    # Rewrite the placeholder paragraph as the new lead-in
    placeholder.text = (
        "The working prototype exposes three primary interfaces, captured below. "
        "Each screenshot is drawn from the live system running via docker compose "
        "against the full 80-document corpus."
    )

    # Append 3 figure blocks (image + caption + description) after the placeholder.
    # python-docx's add_* methods append to end-of-body; we use the underlying
    # XML addnext() to chain them right after the lead-in paragraph.
    anchor_xml = placeholder._element

    for name, caption, desc in SCREENSHOT_SPECS:
        png_path = screenshots[name]

        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        run.add_picture(str(png_path), width=Inches(6.0))

        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_para.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(10)

        desc_para = doc.add_paragraph(desc)
        desc_para.paragraph_format.space_after = Pt(12)

        # Move the three paragraphs from end-of-doc to right after the anchor,
        # chaining so the next iteration lands after the previous description.
        for xml_para in (img_para._element, cap_para._element, desc_para._element):
            anchor_xml.addnext(xml_para)
            anchor_xml = xml_para

    doc.save(docx_path)


def verify_1(docx_path: Path) -> bool:
    doc = Document(docx_path)
    body_text = "\n".join(p.text for p in doc.paragraphs)
    if "will be included in the final version" in body_text:
        raise AssertionError("Patch 1 failed: placeholder text still present.")
    if len(doc.inline_shapes) < 14:
        raise AssertionError(
            f"Patch 1 failed: expected >=14 inline shapes, got {len(doc.inline_shapes)}."
        )
    return True


# ═════════════════════════════════════════════════════════════════════
# Patch 2: simplify jargon for executive audience (§8.2 – §8.7)
# ═════════════════════════════════════════════════════════════════════

# Seven rewrites, each keyed by paragraph index. Technical anchor terms
# (PostgreSQL, RRF, HNSW, embeddings) are preserved; business-impact
# clauses are added so an executive reader can follow.
JARGON_REWRITES: dict[int, str] = {
    148: (
        "The Processing Layer contains the system's core business logic. "
        "The search service orchestrates the hybrid retrieval pipeline, "
        "converting each query into a numerical representation, running three "
        "retrieval signals against the database in parallel, and returning a "
        "ranked result list within milliseconds. Reranking and access-control "
        "filtering run as final steps before results reach the user."
    ),
    149: (
        "The Data Layer centers on a single database instance running the "
        "ParadeDB image, which bundles PostgreSQL with two extensions the "
        "system relies on: pgvector for semantic (meaning-based) search and "
        "pg_search for keyword search. Consolidating both capabilities into "
        "one database removes the operational overhead of running a separate "
        "vector store, and keeps every search signal within the same "
        "transactional boundary."
    ),
    150: (
        "External services include the Qwen3-Embedding-0.6B model (deployed "
        "locally or via an external GPU host) and the Cohere Rerank v3.5 API. "
        "The embedding model converts text to a numerical representation the "
        "database can search; the reranker reorders the top candidates by "
        "deeper relevance before they are shown to the user."
    ),
    156: (
        "The system performs three retrieval signals in parallel against a "
        "single database call. A semantic search finds chunks whose meaning "
        "matches the query; a keyword search finds chunks whose words match; "
        "and a title search boosts documents whose title alone closely "
        "matches the query. Running these in parallel keeps total latency "
        "bounded by the slowest single signal, not their sum."
    ),
    157: (
        "The three ranked lists are merged using Reciprocal Rank Fusion (RRF), "
        "a rank aggregation method that combines the strengths of each signal "
        "without requiring hand-tuned weights. RRF rewards documents that "
        "rank highly across multiple signals and penalizes documents that "
        "appear only in one, which makes the final ranking resilient to "
        "any single signal's weaknesses."
    ),
    169: (
        "The document_chunks table stores the searchable representation of "
        "each document. Every row holds a segment of text (roughly 512 "
        "tokens), its numerical embedding, its parent document ID, and the "
        "access-control level it inherits. A specialized index on the "
        "embedding column (HNSW) enables meaning-based lookups in "
        "milliseconds, even as the table grows."
    ),
    183: (
        "The fourth decision was the merge strategy for combining the three "
        "retrieval signals. Rather than a linear combination, which would "
        "require hand-tuned weights and would quietly drift as the corpus "
        "changes, the system uses Reciprocal Rank Fusion. RRF is "
        "parameter-free relative to query content, scales naturally from "
        "dozens to millions of documents, and degrades gracefully if any "
        "one signal underperforms."
    ),
}


def patch_2_simplify_jargon(docx_path: Path) -> None:
    doc = Document(docx_path)
    for idx, new_text in JARGON_REWRITES.items():
        para = doc.paragraphs[idx]
        para.text = new_text
    doc.save(docx_path)


def verify_2(docx_path: Path) -> bool:
    doc = Document(docx_path)
    for idx, new_text in JARGON_REWRITES.items():
        actual = doc.paragraphs[idx].text
        if actual.strip() != new_text.strip():
            raise AssertionError(
                f"Patch 2 failed at paragraph {idx}: expected rewrite not present."
            )
    return True


# ═════════════════════════════════════════════════════════════════════
# Patch 3: normalize table formatting (data tables only, not callouts)
# ═════════════════════════════════════════════════════════════════════

def _set_cell_fill(cell, fill_hex: str) -> None:
    tcPr = cell._element.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _clear_cell_fill(cell) -> None:
    tcPr = cell._element.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)


def _style_header_cell(cell) -> None:
    _set_cell_fill(cell, L.HEADER_FILL_HEX)
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(L.FONT_SIZE_TABLE_PT)
            run.font.name = L.FONT_NAME
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _style_body_cell(cell, shaded: bool) -> None:
    if shaded:
        _set_cell_fill(cell, L.ROW_ALT_FILL_HEX)
    else:
        _clear_cell_fill(cell)
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.size = Pt(L.FONT_SIZE_TABLE_PT)
            run.font.name = L.FONT_NAME


def patch_3_normalize_tables(docx_path: Path) -> None:
    doc = Document(docx_path)
    for t_idx in L.TABLES_DATA:
        tbl = doc.tables[t_idx]
        # Header row
        for cell in tbl.rows[0].cells:
            _style_header_cell(cell)
        # Body rows with alternating shading
        for row_idx in range(1, len(tbl.rows)):
            shaded = (row_idx % 2 == 1)
            for cell in tbl.rows[row_idx].cells:
                _style_body_cell(cell, shaded)
    doc.save(docx_path)


def verify_3(docx_path: Path) -> bool:
    doc = Document(docx_path)
    for t_idx in L.TABLES_DATA:
        tbl = doc.tables[t_idx]
        for cell in tbl.rows[0].cells:
            tcPr = cell._element.find(qn("w:tcPr"))
            shd = tcPr.find(qn("w:shd")) if tcPr is not None else None
            if shd is None or shd.get(qn("w:fill")) != L.HEADER_FILL_HEX:
                raise AssertionError(
                    f"Patch 3 failed: table {t_idx} header missing navy fill"
                )
    return True


def _run_patches(source: Path, target: Path, up_to: int | None, only: int | None) -> int:
    shutil.copy2(source, target)

    screenshots_dir = HERE / "screenshots"
    screenshots = {
        name: screenshots_dir / name
        for name in ("landing.png", "search-results.png", "admin-upload.png")
    }

    patches = [
        (1, lambda: patch_1_8_10_screenshots(target, screenshots), verify_1),
        (2, lambda: patch_2_simplify_jargon(target), verify_2),
        (3, lambda: patch_3_normalize_tables(target), verify_3),
        # 4..5 land in later tasks
    ]

    for n, run_fn, verify_fn in patches:
        if only is not None and only != n:
            continue
        if up_to is not None and n > up_to:
            break
        print(f"Running patch {n}...")
        run_fn()
        verify_fn(target)
        print(f"  ✓ patch {n} verified")

    print(f"Wrote {target}")
    return 0


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--dry-run", action="store_true")
    ap.add_argument("--scan", action="store_true")
    ap.add_argument("--only", type=int, choices=[1, 2, 3, 4, 5])
    ap.add_argument("--up-to", type=int, choices=[1, 2, 3, 4, 5])
    args = ap.parse_args()

    if args.scan:
        scan(SOURCE)
        return 0

    target = OUTPUT
    if args.dry_run:
        target = HERE / ".dry_run.docx"

    rc = _run_patches(SOURCE, target, up_to=args.up_to, only=args.only)

    if args.dry_run and target.exists():
        target.unlink()
    return rc


if __name__ == "__main__":
    sys.exit(main())
