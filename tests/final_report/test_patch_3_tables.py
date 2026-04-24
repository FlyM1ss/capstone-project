# tests/final_report/test_patch_3_tables.py
"""Patch 3: uniform header shading + alternating rows on all data tables."""
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from apply_edits import patch_3_normalize_tables, verify_3
import locators as L


def _cell_fill(cell) -> str | None:
    tcPr = cell._element.find(qn("w:tcPr"))
    if tcPr is None:
        return None
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        return None
    return shd.get(qn("w:fill"))


def test_patch_3_applies_header_and_alt_shading(fresh_docx_copy: Path):
    patch_3_normalize_tables(fresh_docx_copy)
    doc = Document(fresh_docx_copy)

    for t_idx in L.TABLES_DATA:
        tbl = doc.tables[t_idx]
        # Header row: every cell has fill = HEADER_FILL_HEX
        for c in tbl.rows[0].cells:
            assert _cell_fill(c) == L.HEADER_FILL_HEX, (
                f"Table {t_idx} header cell missing navy fill"
            )
        # Odd data rows (row_idx=1,3,5,...) should have alt fill
        for row_idx in range(1, len(tbl.rows)):
            expected = L.ROW_ALT_FILL_HEX if row_idx % 2 == 1 else None
            for c in tbl.rows[row_idx].cells:
                got = _cell_fill(c)
                if expected is None:
                    assert got in (None, "auto"), (
                        f"Table {t_idx} row {row_idx} unexpected fill {got}"
                    )
                else:
                    assert got == expected, (
                        f"Table {t_idx} row {row_idx} fill got {got} expected {expected}"
                    )

    # Callout tables MUST NOT be touched
    for t_idx in L.TABLES_CALLOUT:
        tbl = doc.tables[t_idx]
        fill = _cell_fill(tbl.rows[0].cells[0])
        assert fill != L.HEADER_FILL_HEX, (
            f"Callout table {t_idx} was incorrectly shaded as a data-table header"
        )

    assert verify_3(fresh_docx_copy) is True
