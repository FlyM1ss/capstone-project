# tests/final_report/test_patch_4_copyedit.py
"""Patch 4: em-dashes removed, redundancies trimmed."""
from pathlib import Path

from docx import Document

from apply_edits import patch_4_copyedit_pass, verify_4


def _all_body_text(doc) -> str:
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return "\n".join(parts)


def test_patch_4_removes_em_dashes(fresh_docx_copy: Path):
    patch_4_copyedit_pass(fresh_docx_copy)
    doc = Document(fresh_docx_copy)
    text = _all_body_text(doc)
    assert "—" not in text, (
        f"Expected 0 em-dashes; found {text.count(chr(0x2014))}"
    )
    assert "in order to" not in text.lower()
    assert verify_4(fresh_docx_copy) is True
